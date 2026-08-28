"""
Builds the proposal as a Word document.

PDF is a dead end for this workflow: procurement teams edit the draft, drop it
on their own letterhead and circulate it with tracked changes. They need .docx.

The drafted response is markdown-ish (the models emit `##` headings, `-`
bullets and `**bold**`), so it's parsed into real Word constructs — built-in
Heading styles and List Bullet/Number styles — rather than dumped as one blob
of preformatted text. That's what makes the file editable and navigable
instead of merely openable.
"""

import io
import re
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# Trim the model's fenced code markers; they're never meaningful in a proposal.
_FENCE = re.compile(r"^\s*```.*$")
_HEADING = re.compile(r"^(#{1,4})\s+(.*)$")
_BULLET = re.compile(r"^\s*[-*•]\s+(.*)$")
_NUMBERED = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_BOLD_SEGMENT = re.compile(r"\*\*(.+?)\*\*")


def _add_rich_text(paragraph, text: str) -> None:
    """Writes `text` into `paragraph`, turning **bold** into real bold runs."""
    position = 0
    for match in _BOLD_SEGMENT.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        paragraph.add_run(match.group(1)).bold = True
        position = match.end()

    remainder = text[position:]
    if remainder:
        paragraph.add_run(remainder)


def _add_body(document: Document, body: str) -> None:
    """
    Renders the markdown-ish draft into headings, lists and paragraphs.

    Soft-wrapped lines are joined back into one paragraph before rendering.
    Emitting a paragraph per source line looks passable on screen but puts a
    hard break mid-sentence in Word, so the text reflows badly the moment
    anyone edits it — and it breaks `**bold**` that spans a wrap.
    """
    buffer: list[str] = []
    in_fence = False

    def flush() -> None:
        if buffer:
            _add_rich_text(document.add_paragraph(), " ".join(buffer).strip())
            buffer.clear()

    for raw_line in (body or "").splitlines():
        line = raw_line.rstrip()

        # Skip fenced blocks entirely, markers and contents alike.
        if _FENCE.match(line):
            flush()
            in_fence = not in_fence
            continue
        if in_fence:
            continue

        if not line.strip():
            flush()
            continue

        heading = _HEADING.match(line)
        if heading:
            flush()
            text = heading.group(2).strip().strip("*")
            if text:
                document.add_heading(text, level=min(len(heading.group(1)), 4))
            continue

        bullet = _BULLET.match(line)
        if bullet:
            flush()
            _add_rich_text(document.add_paragraph(style="List Bullet"), bullet.group(1))
            continue

        numbered = _NUMBERED.match(line)
        if numbered:
            flush()
            _add_rich_text(document.add_paragraph(style="List Number"), numbered.group(1))
            continue

        buffer.append(line.strip())

    flush()


def build_proposal_docx(
    body: str,
    company_name: str = "",
    source_filename: str = "",
) -> bytes:
    """
    Returns the .docx bytes for a proposal.

    Deliberately contains the proposal only. The win score is an internal
    judgement about the bid's chances — including it in the file a client
    receives would be a bad day for somebody.
    """
    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = document.add_heading(
        f"{company_name.strip()} — Proposal" if company_name.strip() else "Proposal",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # A quiet subtitle line: what this responds to, and when it was drafted.
    meta_bits = []
    if source_filename.strip():
        meta_bits.append(f"In response to {source_filename.strip()}")
    meta_bits.append(datetime.now(timezone.utc).strftime("%d %B %Y"))

    subtitle = document.add_paragraph()
    run = subtitle.add_run("  ·  ".join(meta_bits))
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    _add_body(document, body)

    document.core_properties.title = "Proposal"
    if company_name.strip():
        document.core_properties.author = company_name.strip()

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
